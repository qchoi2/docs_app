import hashlib
import json
import sqlite3
from contextlib import closing
from pathlib import Path

from docx import Document

from index_contracts import DEFAULT_ROOT, IndexOptions, build_parser, choose_candidates, index_contracts, main


def short_sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


def write_pdf(path: Path, text: str | None) -> None:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
    ]

    if text is None:
        objects.append(
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << >> >>"
        )
    else:
        stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
        objects.append(
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        )
        objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        objects.append(
            b"<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_at = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        b"trailer\n"
        + f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode("ascii")
        + b"startxref\n"
        + str(xref_at).encode("ascii")
        + b"\n%%EOF\n"
    )
    path.write_bytes(bytes(pdf))


def read_rows(db_path: Path, query: str):
    with closing(sqlite3.connect(db_path)) as conn:
        return conn.execute(query).fetchall()


def write_docx(path: Path, *paragraphs: str) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def test_indexes_synthetic_docx_paragraphs(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()

    docx_path = root / "simple.docx"
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.save(docx_path)

    result = index_contracts(root, out)

    file_key = short_sha256(docx_path.read_bytes())
    txt_path = out / "txt" / f"{file_key}.txt"
    assert result["statuses"] == {"ok": 1}
    assert txt_path.read_text(encoding="utf-8").splitlines() == [
        "[¶1]\tFirst paragraph",
        "[¶2]\tSecond paragraph",
    ]

    rows = read_rows(out / "catalog.sqlite", "SELECT status, path FROM files")
    assert rows == [("ok", "simple.docx")]
    fts_rows = read_rows(
        out / "catalog.sqlite", "SELECT content, para FROM fts ORDER BY para"
    )
    assert fts_rows == [("First paragraph", 1), ("Second paragraph", 2)]


def test_docx_tables_keep_document_order_and_row_paragraphs(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()

    docx_path = root / "table.docx"
    document = Document()
    document.add_paragraph("Before table")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    document.add_paragraph("After table")
    document.save(docx_path)

    index_contracts(root, out)

    file_key = short_sha256(docx_path.read_bytes())
    txt_lines = (out / "txt" / f"{file_key}.txt").read_text(encoding="utf-8").splitlines()
    assert txt_lines == [
        "[¶1]\tBefore table",
        "[¶2]\tA1 | B1",
        "[¶3]\tA2 | B2",
        "[¶4]\tAfter table",
    ]


def test_indexes_text_pdf(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()

    pdf_path = root / "text.pdf"
    write_pdf(pdf_path, "PDF text fixture")

    index_contracts(root, out)

    file_key = short_sha256(pdf_path.read_bytes())
    txt_text = (out / "txt" / f"{file_key}.txt").read_text(encoding="utf-8")
    assert "[¶1]\tPDF text fixture" in txt_text
    rows = read_rows(out / "catalog.sqlite", "SELECT status FROM files")
    assert rows == [("ok",)]
    fts_rows = read_rows(out / "catalog.sqlite", "SELECT file_key FROM fts")
    assert fts_rows == [(file_key,)]


def test_empty_pdf_is_recorded_without_fts_rows(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()

    pdf_path = root / "blank.pdf"
    write_pdf(pdf_path, None)

    index_contracts(root, out)

    file_key = short_sha256(pdf_path.read_bytes())
    assert (out / "txt" / f"{file_key}.txt").read_text(encoding="utf-8") == ""
    file_rows = read_rows(
        out / "catalog.sqlite", "SELECT status, error_reason FROM files"
    )
    assert file_rows == [("empty", "pdf_text_empty")]
    fts_rows = read_rows(out / "catalog.sqlite", "SELECT file_key FROM fts")
    assert fts_rows == []


def test_rerun_same_root_out_skips_unchanged_file(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    write_docx(root / "a.docx", "A")

    first = index_contracts(root, out)
    second = index_contracts(root, out)

    assert first["changes"]["new"] == 1
    assert second["indexed"] == 0
    assert second["skipped"] == 1
    assert second["changes"]["skipped"] == 1


def test_added_file_processes_only_new_file(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    write_docx(root / "a.docx", "A")
    index_contracts(root, out)

    write_docx(root / "b.docx", "B")
    result = index_contracts(root, out)

    assert result["indexed"] == 1
    assert result["skipped"] == 1
    assert result["changes"]["new"] == 1
    rows = read_rows(out / "catalog.sqlite", "SELECT COUNT(*) FROM files")
    assert rows == [(2,)]


def test_deleted_file_becomes_missing_and_fts_is_removed(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    docx_path = root / "a.docx"
    write_docx(docx_path, "A")
    index_contracts(root, out)

    docx_path.unlink()
    result = index_contracts(root, out)

    assert result["changes"]["missing"] == 1
    assert read_rows(out / "catalog.sqlite", "SELECT status FROM files") == [
        ("missing",)
    ]
    assert read_rows(out / "catalog.sqlite", "SELECT COUNT(*) FROM fts") == [(0,)]


def test_file_list_pilot_then_full_expansion_same_root_out(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    write_docx(root / "pilot.docx", "Pilot")
    write_docx(root / "later.docx", "Later")
    file_list = tmp_path / "pilot_files.txt"
    file_list.write_text("pilot.docx\n", encoding="utf-8")

    pilot = index_contracts(root, out, IndexOptions(file_list=file_list))
    full = index_contracts(root, out)

    assert pilot["changes"]["new"] == 1
    assert full["changes"]["new"] == 1
    assert full["skipped"] == 1
    rows = read_rows(
        out / "catalog.sqlite", "SELECT path, status FROM files ORDER BY path"
    )
    assert rows == [("later.docx", "ok"), ("pilot.docx", "ok")]


def test_sample_seed_selection_is_deterministic(tmp_path):
    root = tmp_path / "contracts"
    root.mkdir()
    for index in range(5):
        write_docx(root / f"{index}.docx", str(index))

    first, _ = choose_candidates(root, IndexOptions(sample=3, sample_seed=42))
    second, _ = choose_candidates(root, IndexOptions(sample=3, sample_seed=42))
    third, _ = choose_candidates(root, IndexOptions(sample=3, sample_seed=7))

    assert [path.name for path in first] == [path.name for path in second]
    assert [path.name for path in first] != [path.name for path in third]


def test_file_list_and_sample_together_is_error(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    file_list = tmp_path / "files.txt"
    file_list.write_text("", encoding="utf-8")

    rc = main(
        [
            "--root",
            str(root),
            "--out",
            str(out),
            "--file-list",
            str(file_list),
            "--sample",
            "1",
        ]
    )

    assert rc == 1


def test_root_defaults_to_repo_root_folder():
    args = build_parser().parse_args(["--out", "cs_index"])

    assert args.root == DEFAULT_ROOT


def test_moved_file_updates_path_without_missing(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    original = root / "old.docx"
    moved = root / "new.docx"
    write_docx(original, "Move me")
    file_key = short_sha256(original.read_bytes())
    index_contracts(root, out)

    original.rename(moved)
    result = index_contracts(root, out)

    assert result["changes"]["moved"] == 1
    assert result["changes"]["missing"] == 0
    assert read_rows(out / "catalog.sqlite", "SELECT file_key, path, status FROM files") == [
        (file_key, "new.docx", "ok")
    ]


def test_missing_file_restores_when_same_file_reappears(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    docx_path = root / "a.docx"
    write_docx(docx_path, "Restore")
    original_bytes = docx_path.read_bytes()
    index_contracts(root, out)

    docx_path.unlink()
    index_contracts(root, out)
    docx_path.write_bytes(original_bytes)
    result = index_contracts(root, out)

    assert result["changes"]["restored"] == 1
    assert read_rows(out / "catalog.sqlite", "SELECT status FROM files") == [("ok",)]
    assert read_rows(out / "catalog.sqlite", "SELECT COUNT(*) FROM fts") == [(1,)]


def test_content_change_marks_old_key_missing_and_adds_new_key(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    docx_path = root / "a.docx"
    write_docx(docx_path, "Old")
    old_key = short_sha256(docx_path.read_bytes())
    index_contracts(root, out)

    write_docx(docx_path, "New")
    new_key = short_sha256(docx_path.read_bytes())
    result = index_contracts(root, out)

    assert old_key != new_key
    assert result["changes"]["content_changed"] == 1
    rows = read_rows(
        out / "catalog.sqlite",
        "SELECT file_key, status FROM files ORDER BY status, file_key",
    )
    assert sorted(rows) == sorted([(old_key, "missing"), (new_key, "ok")])


def test_dry_run_writes_report_without_db_or_txt_cache(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    write_docx(root / "a.docx", "Dry")

    result = index_contracts(root, out, IndexOptions(dry_run=True))

    assert result["dry_run"] is True
    assert result["changes"]["new"] == 1
    assert Path(result["report"]).exists()
    assert not (out / "catalog.sqlite").exists()
    assert not (out / "txt").exists()


def test_full_reprocesses_unchanged_file_and_records_batch_label(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    write_docx(root / "a.docx", "A")
    index_contracts(root, out, IndexOptions(batch_label="pilot_001"))

    result = index_contracts(root, out, IndexOptions(full=True, batch_label="full_001"))

    assert result["full"] is True
    assert result["indexed"] == 1
    assert result["skipped"] == 0
    assert read_rows(out / "catalog.sqlite", "SELECT batch_label FROM files") == [
        ("full_001",)
    ]


def test_include_misc_controls_misc_folder_scan(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    misc = root / "Sample_Docs"
    misc.mkdir(parents=True)
    write_docx(misc / "misc.docx", "Misc")

    excluded = index_contracts(root, out)
    included = index_contracts(root, out, IndexOptions(include_misc=True))

    assert excluded["indexed"] == 0
    assert included["indexed"] == 1
    assert read_rows(out / "catalog.sqlite", "SELECT path FROM files") == [
        ("Sample_Docs/misc.docx",)
    ]


def test_unsupported_files_are_recorded_and_zip_is_report_excluded(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    (root / "legacy.doc").write_bytes(b"legacy")
    (root / "archive.zip").write_bytes(b"PK\x05\x06")

    result = index_contracts(root, out)

    assert result["changes"]["unsupported"] == 1
    assert read_rows(
        out / "catalog.sqlite",
        "SELECT path, status, error_reason FROM files",
    ) == [("legacy.doc", "unsupported", "unsupported_ext")]
    report = Path(result["report"]).read_text(encoding="utf-8")
    assert "- archive.zip" in report
    assert read_rows(out / "catalog.sqlite", "SELECT COUNT(*) FROM fts") == [(0,)]


def test_converted_doc_manifest_indexes_original_doc_path(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    original = root / "legacy.doc"
    original.write_bytes(bytes.fromhex("D0CF11E0") + b"legacy")
    source_sha256 = hashlib.sha256(original.read_bytes()).hexdigest()
    converted = out / "converted" / ("%s.docx" % source_sha256)
    converted.parent.mkdir(parents=True)
    write_docx(converted, "Converted body")
    manifest = {
        "schema_version": 1,
        "items": {
            "legacy.doc": {
                "source_path": "legacy.doc",
                "source_sha256": source_sha256,
                "converted_docx": "converted/%s.docx" % source_sha256,
                "converter_version": "mock-word",
                "status": "ok",
                "error_reason": None,
            }
        },
    }
    (out / "converted" / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False),
        encoding="utf-8",
    )

    result = index_contracts(root, out)

    assert result["changes"]["unsupported"] == 0
    rows = read_rows(
        out / "catalog.sqlite",
        "SELECT path, ext, status, error_reason, source_signals FROM files",
    )
    assert rows[0][0:4] == ("legacy.doc", ".doc", "ok", None)
    signals = json.loads(rows[0][4])
    assert signals["source_format"] == "doc_converted"
    assert signals["converter_version"] == "mock-word"
    assert read_rows(out / "catalog.sqlite", "SELECT content FROM fts") == [
        ("Converted body",)
    ]


def test_arbitrary_unsupported_extension_is_recorded(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    (root / "spreadsheet.xlsx").write_bytes(b"not a contract")

    result = index_contracts(root, out)

    assert result["changes"]["unsupported"] == 1
    assert read_rows(
        out / "catalog.sqlite",
        "SELECT path, status, error_reason FROM files",
    ) == [("spreadsheet.xlsx", "unsupported", "unsupported_ext")]


def test_report_filename_collision_uses_numeric_suffix(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    write_docx(root / "a.docx", "A")

    first = index_contracts(root, out)
    second = index_contracts(root, out)

    assert Path(first["report"]).name.startswith("report_")
    assert Path(second["report"]).name.endswith("-2.md")
    assert Path(first["report"]).exists()
    assert Path(second["report"]).exists()


def test_type_rules_classify_and_report_unclassified_folders(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    spa_dir = root / "SPA_국문"
    unknown_dir = root / "unknown"
    spa_dir.mkdir(parents=True)
    unknown_dir.mkdir()
    write_docx(spa_dir / "final.docx", "계약 본문")
    write_docx(unknown_dir / "mystery.docx", "본문")

    result = index_contracts(root, out)

    rows = read_rows(
        out / "catalog.sqlite",
        "SELECT path, ctype, lang FROM files ORDER BY path",
    )
    assert rows == [
        ("SPA_국문/final.docx", "SPA", "국문"),
        ("unknown/mystery.docx", "미분류", "국문"),
    ]
    report = Path(result["report"]).read_text(encoding="utf-8")
    assert "## 2. Type x Language" in report
    assert "## 3. Unclassified Folders" in report
    assert "- unknown: 1" in report


def test_ctype_classification_uses_path_not_body_text(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    write_docx(root / "Project_SPA.docx", "본문에 신주인수권부사채라는 말이 있어도 파일명은 SPA")

    index_contracts(root, out)

    assert read_rows(out / "catalog.sqlite", "SELECT ctype FROM files") == [("SPA",)]


def test_manual_override_by_path_glob(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    (root / "SPA").mkdir(parents=True)
    write_docx(root / "SPA" / "deal.docx", "계약 본문")
    (tmp_path / "data").mkdir()
    (tmp_path / "data" / "manual_overrides.yaml").write_text(
        'paths:\n  "**/SPA/**":\n    ctype: SPA\n    is_draft: false\n',
        encoding="utf-8",
    )

    index_contracts(root, out)

    rows = read_rows(out / "catalog.sqlite", "SELECT ctype, is_draft FROM files")
    assert rows == [("SPA", 0)]


def test_manual_override_by_file_key(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    docx_path = root / "mystery.docx"
    write_docx(docx_path, "본문")
    file_key = short_sha256(docx_path.read_bytes())
    (tmp_path / "data").mkdir()
    (tmp_path / "data" / "manual_overrides.yaml").write_text(
        "files:\n"
        f'  "{file_key}":\n'
        "    ctype: SPA\n"
        "    lang: 영문\n"
        "    is_draft: true\n"
        "    version_hint: final\n"
        "    file_key: hacked\n"
        "    content_hash: hacked\n",
        encoding="utf-8",
    )

    index_contracts(root, out)

    rows = read_rows(
        out / "catalog.sqlite",
        "SELECT file_key, ctype, lang, is_draft, version_hint, content_hash FROM files",
    )
    assert len(rows) == 1
    key, ctype, lang, is_draft, version_hint, content_hash = rows[0]
    # file_key/content_hash must never be overridden
    assert key == file_key
    assert content_hash != "hacked"
    assert (ctype, lang, is_draft, version_hint) == ("SPA", "영문", 1, "final")


def test_manual_override_priority_auto_then_path_then_file_key(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    (root / "SHA_folder").mkdir(parents=True)
    docx_path = root / "SHA_folder" / "deal.docx"
    write_docx(docx_path, "본문")
    file_key = short_sha256(docx_path.read_bytes())
    (tmp_path / "data").mkdir()
    # auto classification says MOU (type_rules), path override says SHA, file_key says SPA
    (tmp_path / "data" / "type_rules.yaml").write_text(
        'ctype_rules:\n  - ctype: MOU\n    patterns: ["SHA_folder"]\n',
        encoding="utf-8",
    )
    (tmp_path / "data" / "manual_overrides.yaml").write_text(
        "paths:\n"
        '  "**/SHA_folder/**":\n'
        "    ctype: SHA\n"
        "    lang: 국문\n"
        "files:\n"
        f'  "{file_key}":\n'
        "    ctype: SPA\n",
        encoding="utf-8",
    )

    index_contracts(root, out)

    rows = read_rows(out / "catalog.sqlite", "SELECT ctype, lang FROM files")
    # file_key override wins over path override; path override's lang still applies
    assert rows == [("SPA", "국문")]


def test_documents_without_text_do_not_share_dup_group(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    # two unrelated scanned PDFs: both extract to empty text
    write_pdf(root / "scan_a.pdf", None)
    write_pdf(root / "scan_b.pdf", None)
    # make the two files byte-distinct (comment after %%EOF is ignored by parsers)
    with (root / "scan_b.pdf").open("ab") as handle:
        handle.write(b"\n% second scan\n")

    index_contracts(root, out)

    rows = read_rows(
        out / "catalog.sqlite",
        "SELECT file_key, dup_group, status FROM files ORDER BY path",
    )
    assert all(status == "empty" for _, _, status in rows)
    # each empty document keeps its own dup_group (no spurious duplicates)
    assert rows[0][1] == rows[0][0]
    assert rows[1][1] == rows[1][0]
    assert rows[0][1] != rows[1][1]


def test_lang_classification_ignores_body_language_clause(tmp_path):
    root = tmp_path / "contracts"
    out = tmp_path / "cs_index"
    root.mkdir()
    # Korean contract whose body mentions "영문" / "English" in a language clause
    write_docx(
        root / "국내회사_주식매매계약.docx",
        "본 계약은 국문으로 작성되며 영문(English) 번역본과 차이가 있는 경우 국문이 우선한다.",
        "매도인은 매수인에게 주식을 양도한다.",
    )

    index_contracts(root, out)

    rows = read_rows(out / "catalog.sqlite", "SELECT lang FROM files")
    assert rows == [("국문",)]
