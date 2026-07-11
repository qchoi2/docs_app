"""webapp의 job/settings write 계층 테스트 (UI-0.2/0.3 백엔드)."""

import io
import json
import time
from pathlib import Path

import pytest
from docx import Document

import webapp
from lib.catalog import initialize_catalog


def make_index(out: Path) -> None:
    initialize_catalog(out / "catalog.sqlite")


def make_app(tmp_path) -> webapp.App:
    out = tmp_path / "cs_index"
    out.mkdir()
    make_index(out)
    return webapp.App(out)


def call(app, method, path, body=None, query=""):
    payload = json.dumps(body).encode() if body is not None else b""
    environ = {
        "REQUEST_METHOD": method,
        "PATH_INFO": path,
        "QUERY_STRING": query,
        "CONTENT_LENGTH": str(len(payload)),
        "wsgi.input": io.BytesIO(payload),
    }
    captured = {}

    def start_response(status, headers):
        captured["status"] = status
        captured["headers"] = headers

    chunks = app(environ, start_response)
    data = b"".join(chunks)
    status_code = int(captured["status"].split()[0])
    parsed = json.loads(data) if data else None
    return status_code, parsed


def wait_for(app, job_id, statuses, timeout=8.0):
    deadline = time.time() + timeout
    while time.time() < deadline:
        code, job = call(app, "GET", f"/api/jobs/{job_id}")
        if code == 200 and job["status"] in statuses:
            return job
        time.sleep(0.05)
    return call(app, "GET", f"/api/jobs/{job_id}")[1]


def test_root_path_validate_reports_counts(tmp_path):
    app = make_app(tmp_path)
    try:
        root = tmp_path / "contracts"
        (root / "sub").mkdir(parents=True)
        Document().save(root / "a.docx")
        Document().save(root / "sub" / "b.docx")
        (root / "notes.txt").write_text("x", encoding="utf-8")

        code, info = call(app, "POST", "/api/settings/root-path/validate",
                          {"path": str(root)})
        assert code == 200
        assert info["exists"] and info["is_dir"] and info["readable"]
        assert info["supported_file_count"] == 2
        assert info["file_count"] == 3
        assert info["network_drive"] is False

        code, info = call(app, "POST", "/api/settings/root-path/validate",
                          {"path": str(tmp_path / "missing")})
        assert code == 200 and info["exists"] is False

        code, _ = call(app, "POST", "/api/settings/root-path/validate", {"path": ""})
        assert code == 400
    finally:
        app.shutdown()


def test_index_job_runs_and_reports_progress(tmp_path):
    app = make_app(tmp_path)
    try:
        root = tmp_path / "contracts"
        root.mkdir()
        for i in range(3):
            doc = Document()
            doc.add_paragraph(f"손해배상 조항 {i}")
            doc.save(root / f"c{i}.docx")

        code, resp = call(app, "POST", "/api/jobs/index", {"root": str(root)})
        assert code == 202 and resp["status"] == "queued"
        job_id = resp["job_id"]

        job = wait_for(app, job_id, {"completed", "failed"})
        assert job["status"] == "completed", job
        assert job["progress_total"] == 3
        assert '"indexed": 3' in (job["message"] or "")

        # job이 목록에 나타난다
        code, listing = call(app, "GET", "/api/jobs")
        assert code == 200 and any(j["id"] == job_id for j in listing["jobs"])
    finally:
        app.shutdown()


def test_index_job_missing_root_fails_with_code(tmp_path):
    app = make_app(tmp_path)
    try:
        code, resp = call(app, "POST", "/api/jobs/index",
                          {"root": str(tmp_path / "nope")})
        assert code == 202
        job = wait_for(app, resp["job_id"], {"failed", "completed"})
        assert job["status"] == "failed"
        assert job["error_code"] == "ROOT_NOT_FOUND"
    finally:
        app.shutdown()


def test_job_log_records_lifecycle(tmp_path):
    app = make_app(tmp_path)
    try:
        root = tmp_path / "contracts"
        root.mkdir()
        doc = Document()
        doc.add_paragraph("손해배상")
        doc.save(root / "a.docx")

        code, resp = call(app, "POST", "/api/jobs/index", {"root": str(root)})
        job_id = resp["job_id"]
        wait_for(app, job_id, {"completed", "failed"})

        code, log = call(app, "GET", f"/api/jobs/{job_id}/log")
        assert code == 200
        lines = [e["line"] for e in log["entries"]]
        assert "started" in lines
        assert any(l.startswith("completed") for l in lines)

        assert call(app, "GET", "/api/jobs/not-hex/log")[0] == 400
        assert call(app, "GET", "/api/jobs/" + "a" * 32 + "/log")[0] == 404
    finally:
        app.shutdown()


def test_job_get_validation_and_not_found(tmp_path):
    app = make_app(tmp_path)
    try:
        code, _ = call(app, "GET", "/api/jobs/not-hex")
        assert code == 400
        code, _ = call(app, "GET", "/api/jobs/" + "a" * 32)
        assert code == 404
    finally:
        app.shutdown()


def test_write_endpoints_do_not_touch_catalog_user_tables(tmp_path):
    # jobs는 jobs.sqlite에만 있고 catalog.sqlite에는 없어야 한다.
    app = make_app(tmp_path)
    try:
        import sqlite3
        with sqlite3.connect(app.out / "catalog.sqlite") as conn:
            tables = {r[0] for r in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'")}
        assert "jobs" not in tables
        assert (app.out / "jobs.sqlite").exists()
    finally:
        app.shutdown()
