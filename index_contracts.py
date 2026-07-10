from __future__ import annotations

import argparse
import fnmatch
import hashlib
import json
import random
import re
import sqlite3
import sys
import unicodedata
from contextlib import closing
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
from zipfile import BadZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.opc.exceptions import PackageNotFoundError
from pdfminer.high_level import extract_text
from pdfminer.pdfdocument import PDFEncryptionError
from pdfminer.pdfparser import PDFSyntaxError
import yaml

from lib.catalog import CatalogError, initialize_catalog
from lib.console import configure_utf8_stdio
from lib.normalize import normalize


SUPPORTED_EXTENSIONS = {".docx", ".pdf"}
ZIP_EXTENSIONS = {".zip"}
TYPE_RULE_PATHS = (Path("data/type_rules.yaml"), Path(".docs/type_rules.yaml"))
MANUAL_OVERRIDE_PATHS = (
    Path("data/manual_overrides.yaml"),
    Path(".docs/manual_overrides.yaml"),
)
ALLOWED_OVERRIDE_KEYS = ("ctype", "lang", "is_draft", "version_hint")
DEFAULT_ROOT = Path(__file__).resolve().parent / "root"


@dataclass
class ExtractedDocument:
    paragraphs: List[str]
    status: str = "ok"
    error_reason: Optional[str] = None
    warnings: List[str] = field(default_factory=list)


@dataclass
class ExistingRecord:
    file_key: str
    path: str
    size: Optional[int]
    mtime: Optional[float]
    status: str
    content_hash: Optional[str]
    txt_path: Optional[str]


@dataclass
class IndexOptions:
    include_misc: bool = False
    full: bool = False
    batch_label: Optional[str] = None
    file_list: Optional[Path] = None
    sample: Optional[int] = None
    sample_seed: int = 0
    dry_run: bool = False


@dataclass
class ChangeReport:
    new: List[str] = field(default_factory=list)
    skipped: List[str] = field(default_factory=list)
    moved: List[str] = field(default_factory=list)
    missing: List[str] = field(default_factory=list)
    restored: List[str] = field(default_factory=list)
    content_changed: List[str] = field(default_factory=list)
    excluded: List[str] = field(default_factory=list)
    unsupported: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


@dataclass
class TypeRule:
    value: str
    patterns: List[str]
    flag: Optional[str] = None


@dataclass
class TypeRules:
    ctype_rules: List[TypeRule] = field(default_factory=list)
    lang_rules: List[TypeRule] = field(default_factory=list)
    draft_patterns: List[str] = field(default_factory=list)
    executed_patterns: List[str] = field(default_factory=list)
    version_capture: List[str] = field(default_factory=list)


def sha256_short(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


def hash_text(text: str) -> str:
    return sha256_short(text.encode("utf-8"))


def _find_runtime_file(candidates: Tuple[Path, ...], start: Optional[Path] = None) -> Optional[Path]:
    """Look under cwd (or explicit start) first, then next to this script."""
    bases = []
    if start is not None:
        bases.append(start)
    bases.extend([Path.cwd(), Path(__file__).resolve().parent])
    for base in bases:
        for candidate in candidates:
            path = base / candidate
            if path.exists():
                return path
    return None


def load_type_rules(start: Optional[Path] = None) -> TypeRules:
    selected = _find_runtime_file(TYPE_RULE_PATHS, start)
    if selected is None:
        return TypeRules()

    data = yaml.safe_load(selected.read_text(encoding="utf-8")) or {}
    ctype_rules = [
        TypeRule(
            value=str(item.get("ctype", "")),
            patterns=[str(pattern) for pattern in item.get("patterns", [])],
            flag=item.get("flag"),
        )
        for item in data.get("ctype_rules", [])
    ]
    lang_rules = [
        TypeRule(
            value=str(item.get("lang", "")),
            patterns=[str(pattern) for pattern in item.get("patterns", [])],
            flag=item.get("flag"),
        )
        for item in data.get("lang_rules", [])
    ]
    version_rules = data.get("version_rules", {}) or {}
    return TypeRules(
        ctype_rules=ctype_rules,
        lang_rules=lang_rules,
        draft_patterns=[str(pattern) for pattern in version_rules.get("draft_patterns", [])],
        executed_patterns=[str(pattern) for pattern in version_rules.get("executed_patterns", [])],
        version_capture=[str(pattern) for pattern in version_rules.get("version_capture", [])],
    )


@dataclass
class ManualOverrides:
    paths: List[Tuple[str, Dict[str, object]]] = field(default_factory=list)
    files: Dict[str, Dict[str, object]] = field(default_factory=dict)


def _clean_override_values(raw: Optional[Dict[str, object]]) -> Dict[str, object]:
    """Keep only ctype/lang/is_draft/version_hint; file_key/content_hash are never overridden."""
    values: Dict[str, object] = {}
    for key, value in (raw or {}).items():
        if key not in ALLOWED_OVERRIDE_KEYS or value is None:
            continue
        if key == "is_draft":
            values[key] = 1 if bool(value) else 0
        else:
            values[key] = str(value)
    return values


def load_manual_overrides(start: Optional[Path] = None) -> ManualOverrides:
    selected = _find_runtime_file(MANUAL_OVERRIDE_PATHS, start)
    if selected is None:
        return ManualOverrides()

    data = yaml.safe_load(selected.read_text(encoding="utf-8")) or {}
    paths = [
        (str(pattern), _clean_override_values(values))
        for pattern, values in (data.get("paths") or {}).items()
    ]
    files = {
        str(file_key): _clean_override_values(values)
        for file_key, values in (data.get("files") or {}).items()
    }
    return ManualOverrides(paths=paths, files=files)


def _path_pattern_matches(rel_path: str, pattern: str) -> bool:
    rel_path = unicodedata.normalize("NFC", rel_path)
    pattern = unicodedata.normalize("NFC", pattern)
    # "./" prefix lets root-level entries match "**/..." style patterns too.
    return fnmatch.fnmatch(rel_path, pattern) or fnmatch.fnmatch(f"./{rel_path}", pattern)


def apply_manual_overrides(
    rel_path: str,
    file_key: str,
    ctype: str,
    lang: str,
    is_draft: Optional[int],
    version_hint: Optional[str],
    source_signals: str,
    overrides: ManualOverrides,
) -> Tuple[str, str, Optional[int], Optional[str], str]:
    """Apply overrides with priority: auto classification -> path glob -> file_key."""
    applied: List[Dict[str, object]] = []
    values: Dict[str, object] = {}
    for pattern, pattern_values in overrides.paths:
        if pattern_values and _path_pattern_matches(rel_path, pattern):
            values.update(pattern_values)
            applied.append({"source": "path", "pattern": pattern, "values": pattern_values})
    file_values = overrides.files.get(file_key)
    if file_values:
        values.update(file_values)
        applied.append({"source": "file_key", "values": file_values})
    if not values:
        return ctype, lang, is_draft, version_hint, source_signals

    ctype = str(values.get("ctype", ctype))
    lang = str(values.get("lang", lang))
    if "is_draft" in values:
        is_draft = int(values["is_draft"])
    if "version_hint" in values:
        version_hint = str(values["version_hint"])
    try:
        signals = json.loads(source_signals) if source_signals else {}
    except json.JSONDecodeError:
        signals = {}
    if not isinstance(signals, dict):
        signals = {}
    signals["manual_overrides"] = applied
    return ctype, lang, is_draft, version_hint, json.dumps(signals, ensure_ascii=False)


def matches_any(text: str, patterns: List[str]) -> bool:
    # NFC-normalize both sides: macOS-origin paths can arrive as NFD (brief §4).
    lowered = unicodedata.normalize("NFC", text).lower()
    return any(unicodedata.normalize("NFC", pattern).lower() in lowered for pattern in patterns)


def classify_path(rel_path: str, text: str, rules: TypeRules) -> Tuple[str, str, Optional[int], Optional[str], str]:
    path_haystack = rel_path
    lang_haystack = f"{rel_path} {text}"
    ctype = "미분류"
    lang = "미상"
    source_signals = {}

    for rule in rules.ctype_rules:
        if matches_any(path_haystack, rule.patterns):
            ctype = rule.value
            source_signals["ctype"] = {"value": ctype, "patterns": rule.patterns}
            break

    for rule in rules.lang_rules:
        if matches_any(lang_haystack, rule.patterns):
            lang = rule.value
            source_signals["lang"] = {"value": lang, "patterns": rule.patterns}
            break

    if lang == "미상" and text:
        korean_chars = len(re.findall(r"[가-힣]", text))
        ascii_letters = len(re.findall(r"[A-Za-z]", text))
        total = max(korean_chars + ascii_letters, 1)
        korean_ratio = korean_chars / total
        if korean_ratio > 0.30:
            lang = "국문"
        elif korean_ratio < 0.05 and ascii_letters:
            lang = "영문"
        else:
            lang = "국영문"
        source_signals["lang_heuristic"] = {"korean_ratio": round(korean_ratio, 3)}

    name = Path(rel_path).name
    is_draft = None
    if matches_any(name, rules.draft_patterns):
        is_draft = 1
    elif matches_any(name, rules.executed_patterns):
        is_draft = 0

    version_hint = None
    for pattern in rules.version_capture:
        try:
            match = re.search(pattern, name, re.IGNORECASE)
        except re.error:
            continue
        if match:
            version_hint = match.group(0)
            break

    return ctype, lang, is_draft, version_hint, json.dumps(source_signals, ensure_ascii=False)


def is_misc_path(path: Path, rules: TypeRules) -> bool:
    path_text = path.as_posix()
    return any(
        rule.flag == "exclude_by_default" and matches_any(path_text, rule.patterns)
        for rule in rules.ctype_rules
    )


def is_supported_file(path: Path) -> bool:
    return path.is_file() and not path.is_symlink() and path.suffix.lower() in SUPPORTED_EXTENSIONS


def is_indexable_file(path: Path) -> bool:
    return path.is_file() and not path.is_symlink()


def iter_indexable_files(root: Path, include_misc: bool, rules: TypeRules) -> List[Path]:
    return sorted(
        path
        for path in root.rglob("*")
        if is_indexable_file(path)
        and (include_misc or not is_misc_path(path.relative_to(root), rules))
    )


def read_file_list(root: Path, file_list: Path, include_misc: bool, rules: TypeRules) -> List[Path]:
    paths: List[Path] = []
    for raw_line in file_list.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue
        path = (root / line).resolve()
        try:
            rel_path = path.relative_to(root)
        except ValueError as exc:
            raise ValueError(f"file-list entry escapes root: {line}") from exc
        if not path.exists() or not path.is_file() or path.is_symlink():
            raise ValueError(f"file-list entry is not an existing file: {line}")
        if not is_indexable_file(path):
            continue
        if include_misc or not is_misc_path(rel_path, rules):
            paths.append(path)
    return sorted(set(paths))


def choose_candidates(root: Path, options: IndexOptions, rules: Optional[TypeRules] = None) -> Tuple[List[Path], bool]:
    rules = rules or load_type_rules()
    if options.file_list and options.sample is not None:
        raise ValueError("--file-list and --sample cannot be used together")

    if options.file_list:
        return read_file_list(root, options.file_list, options.include_misc, rules), False

    files = iter_indexable_files(root, options.include_misc, rules)
    if options.sample is not None:
        if options.sample < 0:
            raise ValueError("--sample must be >= 0")
        rng = random.Random(options.sample_seed)
        sample_pool = [path for path in files if path.suffix.lower() in SUPPORTED_EXTENSIONS]
        selected = rng.sample(sample_pool, min(options.sample, len(sample_pool)))
        return sorted(selected), False

    return files, True


def iter_docx_body_blocks(document: DocumentObject):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def normalize_paragraphs(paragraphs: List[str]) -> List[str]:
    normalized: List[str] = []
    for paragraph in paragraphs:
        text = normalize(paragraph)
        if text:
            normalized.append(text)
    return normalized


def extract_docx(path: Path) -> ExtractedDocument:
    try:
        document = Document(path)
    except (BadZipFile, PackageNotFoundError) as exc:
        return ExtractedDocument([], "error", "corrupt_docx", [str(exc)])
    except Exception as exc:
        return ExtractedDocument([], "error", "docx_extract_failed", [str(exc)])

    paragraphs: List[str] = []
    warnings: List[str] = []

    try:
        for block in iter_docx_body_blocks(document):
            if isinstance(block, Paragraph):
                paragraphs.append(block.text)
            elif isinstance(block, Table):
                for row in block.rows:
                    cell_text = [normalize(cell.text) for cell in row.cells]
                    paragraphs.append(" | ".join(text for text in cell_text if text))
    except Exception as exc:
        return ExtractedDocument([], "error", "docx_extract_failed", [str(exc)])

    try:
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                if normalize(paragraph.text):
                    paragraphs.append(f"[머리글] {paragraph.text}")
            for paragraph in section.footer.paragraphs:
                if normalize(paragraph.text):
                    paragraphs.append(f"[바닥글] {paragraph.text}")
    except Exception as exc:
        warnings.append(f"header_footer_extract_skipped: {exc}")

    warnings.append("footnote_extract_skipped")
    normalized = normalize_paragraphs(paragraphs)
    if not normalized:
        return ExtractedDocument([], "empty", "empty_text", warnings)
    return ExtractedDocument(normalized, "ok", None, warnings)


def extract_pdf(path: Path) -> ExtractedDocument:
    try:
        text = extract_text(path)
    except PDFEncryptionError as exc:
        return ExtractedDocument([], "error", "encrypted_pdf", [str(exc)])
    except (PDFSyntaxError, Exception) as exc:
        return ExtractedDocument([], "error", "pdf_extract_failed", [str(exc)])

    paragraphs = normalize_paragraphs(text.split("\n\n"))
    if not paragraphs:
        return ExtractedDocument([], "empty", "pdf_text_empty")
    return ExtractedDocument(paragraphs)


def extract_file(path: Path) -> ExtractedDocument:
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".pdf":
        return extract_pdf(path)
    return ExtractedDocument([], "unsupported", "unsupported_ext")


def error_reason_for_os_error(exc: OSError) -> str:
    if isinstance(exc, PermissionError):
        return "permission_denied"
    return "unknown_error"


def write_txt_cache(out_dir: Path, file_key: str, paragraphs: List[str]) -> Path:
    txt_dir = out_dir / "txt"
    txt_dir.mkdir(parents=True, exist_ok=True)
    txt_path = txt_dir / f"{file_key}.txt"
    lines = [f"[¶{index}]\t{text}" for index, text in enumerate(paragraphs, start=1)]
    txt_path.write_text("\n".join(lines) + ("\n" if lines else ""), encoding="utf-8")
    return txt_path


def record_file(
    conn: sqlite3.Connection,
    *,
    root: Path,
    path: Path,
    out_dir: Path,
    file_key: str,
    content_hash: str,
    extracted: ExtractedDocument,
    txt_path: Path,
    batch_label: Optional[str] = None,
    ctype: str = "미분류",
    lang: str = "미상",
    source_signals: str = "{}",
    is_draft: Optional[int] = None,
    version_hint: Optional[str] = None,
) -> None:
    stat = path.stat()
    rel_path = path.relative_to(root).as_posix()
    text_for_hash = "\n".join(extracted.paragraphs)
    now = datetime.now(timezone.utc).isoformat()

    conn.execute(
        """
        INSERT OR REPLACE INTO files (
          file_key, path, folder, filename, ctype, lang, ext, size, mtime,
          txt_path, char_count, status, error_reason,
          source_signals, batch_label, content_hash, dup_group, is_draft,
          version_hint, indexed_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            file_key,
            rel_path,
            Path(rel_path).parent.as_posix() if Path(rel_path).parent.as_posix() != "." else "",
            path.name,
            ctype,
            lang,
            path.suffix.lower(),
            stat.st_size,
            stat.st_mtime,
            txt_path.relative_to(out_dir).as_posix(),
            len(text_for_hash),
            extracted.status,
            extracted.error_reason,
            source_signals,
            batch_label,
            content_hash,
            file_key,
            is_draft,
            version_hint,
            now,
        ),
    )

    conn.execute("DELETE FROM fts WHERE file_key = ?", (file_key,))
    if extracted.status == "ok":
        conn.executemany(
            "INSERT INTO fts(content, file_key, para) VALUES (?, ?, ?)",
            [
                (paragraph, file_key, index)
                for index, paragraph in enumerate(extracted.paragraphs, start=1)
            ],
        )


def load_existing_records(conn: sqlite3.Connection) -> List[ExistingRecord]:
    rows = conn.execute(
        """
        SELECT file_key, path, size, mtime, status, content_hash, txt_path
        FROM files
        """
    ).fetchall()
    return [ExistingRecord(*row) for row in rows]


def mark_missing(conn: sqlite3.Connection, file_key: str) -> None:
    conn.execute(
        "UPDATE files SET status = 'missing', error_reason = NULL WHERE file_key = ?",
        (file_key,),
    )
    conn.execute("DELETE FROM fts WHERE file_key = ?", (file_key,))


def rebuild_dup_groups(conn: sqlite3.Connection) -> None:
    """Group only successfully extracted documents by content hash.

    Documents without extracted text (empty/error/unsupported) all share the
    hash of the empty string, so grouping them by content_hash would join
    unrelated scanned contracts into one spurious dup_group. They keep their
    own file_key as dup_group instead.
    """
    rows = conn.execute(
        "SELECT file_key, content_hash, status FROM files ORDER BY file_key"
    ).fetchall()
    groups: Dict[str, List[str]] = {}
    solo_keys: List[str] = []
    for file_key, content_hash, status in rows:
        if status == "ok" and content_hash:
            groups.setdefault(content_hash, []).append(file_key)
        else:
            solo_keys.append(file_key)
    conn.executemany(
        "UPDATE files SET dup_group = ? WHERE file_key = ?",
        [(key, key) for key in solo_keys],
    )
    for keys in groups.values():
        dup_group = min(keys)
        conn.executemany(
            "UPDATE files SET dup_group = ? WHERE file_key = ?",
            [(dup_group, key) for key in keys],
        )


def summarize_db(conn: sqlite3.Connection) -> Dict[str, object]:
    status_counts = dict(
        conn.execute("SELECT status, COUNT(*) FROM files GROUP BY status").fetchall()
    )
    batch_counts = dict(
        conn.execute(
            "SELECT COALESCE(batch_label, ''), COUNT(*) FROM files GROUP BY batch_label"
        ).fetchall()
    )
    error_counts = dict(
        conn.execute(
            """
            SELECT error_reason, COUNT(*)
            FROM files
            WHERE error_reason IS NOT NULL
            GROUP BY error_reason
            """
        ).fetchall()
    )
    duplicate_groups = conn.execute(
        """
        SELECT dup_group, COUNT(*)
        FROM files
        WHERE dup_group IS NOT NULL AND status != 'missing'
        GROUP BY dup_group
        HAVING COUNT(*) >= 2
        ORDER BY COUNT(*) DESC, dup_group
        """
    ).fetchall()
    existing_duplicate_groups = conn.execute(
        """
        SELECT dup_group, file_key, path
        FROM files
        WHERE dup_group IN (
          SELECT dup_group
          FROM files
          WHERE dup_group IS NOT NULL AND status != 'missing'
          GROUP BY dup_group
          HAVING COUNT(*) >= 2
        )
          AND status != 'missing'
        ORDER BY dup_group, path
        """
    ).fetchall()
    stale_doc_meta = conn.execute(
        """
        SELECT COUNT(*)
        FROM doc_meta dm
        JOIN files f ON f.file_key = dm.file_key
        WHERE dm.txt_hash IS NOT NULL
          AND f.content_hash IS NOT NULL
          AND dm.txt_hash != f.content_hash
        """
    ).fetchone()[0]
    type_lang = conn.execute(
        """
        SELECT ctype, lang, COUNT(*)
        FROM files
        WHERE status != 'missing'
        GROUP BY ctype, lang
        ORDER BY ctype, lang
        """
    ).fetchall()
    unclassified_folders = conn.execute(
        """
        SELECT COALESCE(NULLIF(folder, ''), '.'), COUNT(*)
        FROM files
        WHERE ctype = '미분류' AND status != 'missing'
        GROUP BY COALESCE(NULLIF(folder, ''), '.')
        ORDER BY COUNT(*) DESC, 1
        """
    ).fetchall()
    return {
        "statuses": status_counts,
        "batch_labels": batch_counts,
        "errors": error_counts,
        "duplicate_groups": duplicate_groups,
        "duplicate_group_members": existing_duplicate_groups,
        "stale_doc_meta": stale_doc_meta,
        "type_lang": type_lang,
        "unclassified_folders": unclassified_folders,
    }


def write_report(
    out_dir: Path,
    result: Dict[str, object],
    changes: ChangeReport,
    db_summary: Optional[Dict[str, object]],
) -> Path:
    report_path = unique_report_path(out_dir)
    new_paths = {item.split(": ", 1)[0] for item in changes.new}
    changed_new_keys = {
        item.split("->", 1)[1]
        for item in changes.content_changed
        if "->" in item
    }
    new_duplicate_groups = []
    if db_summary is not None:
        group_members: Dict[str, List[Tuple[str, str]]] = {}
        for dup_group, file_key, path in db_summary["duplicate_group_members"]:
            group_members.setdefault(dup_group, []).append((file_key, path))
        for dup_group, count in db_summary["duplicate_groups"]:
            members = group_members.get(dup_group, [])
            if any(file_key in changed_new_keys or path in new_paths for file_key, path in members):
                new_duplicate_groups.append((dup_group, count))

    lines = [
        "# Index Report",
        "",
        "## 1. Run And Changes",
        "",
        f"- root: `{result['root']}`",
        f"- out: `{result['out']}`",
        f"- dry_run: `{result['dry_run']}`",
        f"- full: `{result['full']}`",
        f"- batch_label: `{result.get('batch_label') or ''}`",
        "",
        "## Changes",
        "",
        f"- new: {len(changes.new)}",
        f"- skipped: {len(changes.skipped)}",
        f"- moved: {len(changes.moved)}",
        f"- missing: {len(changes.missing)}",
        f"- restored: {len(changes.restored)}",
        f"- content_changed: {len(changes.content_changed)}",
        f"- excluded: {len(changes.excluded)}",
        f"- unsupported: {len(changes.unsupported)}",
        f"- errors: {len(changes.errors)}",
        "",
    ]
    for title, items in [
        ("New", changes.new),
        ("Moved", changes.moved),
        ("Missing", changes.missing),
        ("Restored", changes.restored),
        ("Content Changed", changes.content_changed),
    ]:
        lines.extend([f"### {title}", ""])
        if items:
            lines.extend(f"- {item}" for item in items)
        else:
            lines.append("- none")
        lines.append("")

    if db_summary is None:
        lines.extend(["## 2. Type x Language", ""])
        lines.append("- dry-run: database not modified")
        lines.extend(["", "## 3. Unclassified Folders", "", "- dry-run: database not modified"])
        lines.extend(["", "## 4. Status Counts", "", "- dry-run: database not modified"])
        lines.extend(["", "## 5. Duplicate Groups", "", "- dry-run: database not modified"])
        lines.extend(["", "## 6. Unsupported And Excluded", ""])
        lines.extend(f"- {item} (unsupported)" for item in changes.unsupported)
        lines.extend(f"- {item} (excluded)" for item in changes.excluded)
        if not changes.unsupported and not changes.excluded:
            lines.append("- none")
        lines.extend(["", "## 7. Error Reasons", "", "- dry-run: database not modified"])
        lines.extend(["", "## 8. Batch Labels", "", "- dry-run: database not modified"])
        lines.extend(["", "## 9. Stale Doc Meta", "", "- dry-run: database not modified"])
    else:
        lines.extend(["## 2. Type x Language", ""])
        type_lang = db_summary["type_lang"]
        if type_lang:
            lines.extend(f"- {ctype} / {lang}: {count}" for ctype, lang, count in type_lang)
        else:
            lines.append("- none")
        lines.extend(["", "## 3. Unclassified Folders", ""])
        unclassified = db_summary["unclassified_folders"]
        if unclassified:
            lines.extend(f"- {folder}: {count}" for folder, count in unclassified)
        else:
            lines.append("- none")
        lines.extend(["", "## 4. Status Counts", ""])
        if db_summary["statuses"]:
            lines.extend(f"- {status}: {count}" for status, count in db_summary["statuses"].items())
        else:
            lines.append("- none")
        lines.extend(["", "## 5. Duplicate Groups", ""])
        if db_summary["duplicate_groups"]:
            for dup_group, count in db_summary["duplicate_groups"]:
                marker = " (new in this run)" if (dup_group, count) in new_duplicate_groups else ""
                lines.append(f"- {dup_group}: {count}{marker}")
        else:
            lines.append("- none")
        lines.extend(["", "## 6. Unsupported And Excluded", ""])
        if changes.unsupported or changes.excluded:
            lines.extend(f"- {item} (unsupported)" for item in changes.unsupported)
            lines.extend(f"- {item} (excluded)" for item in changes.excluded)
        else:
            lines.append("- none")
        lines.extend(["", "## 7. Error Reasons", ""])
        if db_summary["errors"]:
            lines.extend(f"- {reason}: {count}" for reason, count in db_summary["errors"].items())
        else:
            lines.append("- none")
        lines.extend(["", "## 8. Batch Labels", ""])
        if db_summary["batch_labels"]:
            lines.extend(f"- {label or '(none)'}: {count}" for label, count in db_summary["batch_labels"].items())
        else:
            lines.append("- none")
        lines.extend(["", "## 9. Stale Doc Meta", ""])
        lines.append(f"- stale doc_meta: {db_summary['stale_doc_meta']}")
    lines.append("")

    out_dir.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(lines), encoding="utf-8")
    return report_path


def unique_report_path(out_dir: Path) -> Path:
    stem = f"report_{datetime.now().strftime('%Y%m%d')}"
    first = out_dir / f"{stem}.md"
    if not first.exists():
        return first
    suffix = 2
    while True:
        candidate = out_dir / f"{stem}-{suffix}.md"
        if not candidate.exists():
            return candidate
        suffix += 1


def index_contracts(
    root: Union[str, Path],
    out: Union[str, Path],
    options: Optional[IndexOptions] = None,
) -> Dict[str, object]:
    options = options or IndexOptions()
    root_path = Path(root).resolve()
    out_dir = Path(out).resolve()

    if not root_path.exists() or not root_path.is_dir():
        raise ValueError(f"root must be an existing directory: {root_path}")

    rules = load_type_rules()
    overrides = load_manual_overrides()
    candidates, is_full_scan = choose_candidates(root_path, options, rules)
    db_path = out_dir / "catalog.sqlite"
    if not options.dry_run:
        out_dir.mkdir(parents=True, exist_ok=True)
        db_path = initialize_catalog(db_path)

    indexed = 0
    skipped = 0
    statuses: Dict[str, int] = {}
    warnings: Dict[str, int] = {}
    changes = ChangeReport()
    db_summary: Optional[Dict[str, object]] = None

    if options.dry_run and not db_path.exists():
        existing_records: List[ExistingRecord] = []
        conn = None
    else:
        if options.dry_run:
            conn = sqlite3.connect(db_path)
        else:
            conn = sqlite3.connect(db_path)

    try:
        if conn is not None and options.full and not options.dry_run:
            conn.execute("DELETE FROM fts")
            conn.execute("DELETE FROM files")
            conn.commit()

        existing_records = load_existing_records(conn) if conn is not None else []
        by_path = {record.path: record for record in existing_records}
        by_key = {record.file_key: record for record in existing_records}
        seen_paths = set()
        seen_keys = set()

        for path in candidates:
            rel_path = path.relative_to(root_path).as_posix()
            if path.suffix.lower() in ZIP_EXTENSIONS:
                changes.excluded.append(rel_path)
                continue

            seen_paths.add(rel_path)
            try:
                stat = path.stat()
            except OSError as exc:
                changes.errors.append(f"{rel_path}: {error_reason_for_os_error(exc)}")
                continue

            existing_at_path = by_path.get(rel_path)
            if (
                existing_at_path
                and existing_at_path.status != "missing"
                and existing_at_path.size == stat.st_size
                and existing_at_path.mtime == stat.st_mtime
                and not options.full
            ):
                skipped += 1
                changes.skipped.append(rel_path)
                seen_keys.add(existing_at_path.file_key)
                continue

            try:
                file_bytes = path.read_bytes()
            except OSError as exc:
                reason = error_reason_for_os_error(exc)
                file_key = sha256_short(f"unreadable:{rel_path}".encode("utf-8"))
                extracted = ExtractedDocument([], "error", reason, [str(exc)])
                content_hash = hash_text("")
                txt_path = out_dir / "txt" / f"{file_key}.txt"
                if conn is not None and not options.dry_run:
                    txt_path = write_txt_cache(out_dir, file_key, [])
                    ctype, lang, is_draft, version_hint, source_signals = classify_path(
                        rel_path, "", rules
                    )
                    ctype, lang, is_draft, version_hint, source_signals = apply_manual_overrides(
                        rel_path, file_key, ctype, lang, is_draft, version_hint,
                        source_signals, overrides,
                    )
                    record_file(
                        conn,
                        root=root_path,
                        path=path,
                        out_dir=out_dir,
                        file_key=file_key,
                        content_hash=content_hash,
                        extracted=extracted,
                        txt_path=txt_path,
                        batch_label=options.batch_label,
                        ctype=ctype,
                        lang=lang,
                        source_signals=source_signals,
                        is_draft=is_draft,
                        version_hint=version_hint,
                    )
                indexed += 1
                statuses[extracted.status] = statuses.get(extracted.status, 0) + 1
                changes.errors.append(f"{rel_path}: {reason}")
                continue

            file_key = sha256_short(file_bytes)
            seen_keys.add(file_key)
            existing_same_key = by_key.get(file_key)

            if (
                existing_at_path
                and existing_at_path.file_key != file_key
                and existing_at_path.status != "missing"
            ):
                changes.content_changed.append(
                    f"{rel_path}: {existing_at_path.file_key}->{file_key}"
                )
                if conn is not None and not options.dry_run:
                    mark_missing(conn, existing_at_path.file_key)
            elif existing_same_key and existing_same_key.path != rel_path:
                changes.moved.append(f"{existing_same_key.path}->{rel_path}")
            elif (
                (existing_at_path and existing_at_path.status == "missing")
                or (existing_same_key and existing_same_key.status == "missing")
            ):
                changes.restored.append(rel_path)
            elif not existing_at_path and not existing_same_key:
                changes.new.append(rel_path)

            extracted = extract_file(path)
            text_for_hash = "\n".join(extracted.paragraphs)
            content_hash = hash_text(text_for_hash)
            txt_path = out_dir / "txt" / f"{file_key}.txt"
            ctype, lang, is_draft, version_hint, source_signals = classify_path(
                rel_path, text_for_hash, rules
            )
            ctype, lang, is_draft, version_hint, source_signals = apply_manual_overrides(
                rel_path, file_key, ctype, lang, is_draft, version_hint,
                source_signals, overrides,
            )
            if extracted.status == "unsupported":
                changes.unsupported.append(rel_path)

            if conn is not None and not options.dry_run:
                txt_path = write_txt_cache(out_dir, file_key, extracted.paragraphs)
                record_file(
                    conn,
                    root=root_path,
                    path=path,
                    out_dir=out_dir,
                    file_key=file_key,
                    content_hash=content_hash,
                    extracted=extracted,
                    txt_path=txt_path,
                    batch_label=options.batch_label,
                    ctype=ctype,
                    lang=lang,
                    source_signals=source_signals,
                    is_draft=is_draft,
                    version_hint=version_hint,
                )

            indexed += 1
            statuses[extracted.status] = statuses.get(extracted.status, 0) + 1
            for warning in extracted.warnings:
                key = warning.split(":", 1)[0]
                warnings[key] = warnings.get(key, 0) + 1

        if is_full_scan:
            current_supported_paths = set(seen_paths)
            for record in existing_records:
                if (
                    record.status != "missing"
                    and record.path not in current_supported_paths
                    and record.file_key not in seen_keys
                ):
                    changes.missing.append(record.path)
                    if conn is not None and not options.dry_run:
                        mark_missing(conn, record.file_key)

        if conn is not None and not options.dry_run:
            rebuild_dup_groups(conn)
            conn.commit()
            db_summary = summarize_db(conn)
            conn.execute("PRAGMA wal_checkpoint(TRUNCATE)")
        elif conn is not None:
            db_summary = summarize_db(conn)
    finally:
        if conn is not None:
            conn.close()

    result = {
        "root": str(root_path),
        "out": str(out_dir),
        "db": str(db_path),
        "indexed": indexed,
        "skipped": skipped,
        "dry_run": options.dry_run,
        "full": options.full,
        "batch_label": options.batch_label,
        "statuses": statuses,
        "warnings": warnings,
        "changes": {
            "new": len(changes.new),
            "skipped": len(changes.skipped),
            "moved": len(changes.moved),
            "missing": len(changes.missing),
            "restored": len(changes.restored),
            "content_changed": len(changes.content_changed),
            "excluded": len(changes.excluded),
            "unsupported": len(changes.unsupported),
            "errors": len(changes.errors),
        },
    }
    report_path = write_report(out_dir, result, changes, db_summary)
    result["report"] = str(report_path)
    return result


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Index DOCX/PDF contracts.")
    parser.add_argument(
        "--root",
        default=DEFAULT_ROOT,
        type=Path,
        help="Root folder to scan. Defaults to ./root next to index_contracts.py.",
    )
    parser.add_argument("--out", required=True, help="Output cs_index folder.")
    parser.add_argument("--include-misc", action="store_true")
    parser.add_argument("--full", action="store_true")
    parser.add_argument("--batch-label")
    parser.add_argument("--file-list", type=Path)
    parser.add_argument("--sample", type=int)
    parser.add_argument("--sample-seed", type=int, default=0)
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--quiet", action="store_true", help="Suppress the JSON summary on stdout.")
    return parser


def main(argv: Optional[List[str]] = None) -> int:
    configure_utf8_stdio()
    parser = build_parser()
    args = parser.parse_args(argv)

    try:
        result = index_contracts(
            args.root,
            args.out,
            IndexOptions(
                include_misc=args.include_misc,
                full=args.full,
                batch_label=args.batch_label,
                file_list=args.file_list,
                sample=args.sample,
                sample_seed=args.sample_seed,
                dry_run=args.dry_run,
            ),
        )
    except (CatalogError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    if not args.quiet:
        print(json.dumps(result, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
